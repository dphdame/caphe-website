#!/usr/bin/env python3
"""
Phase 6: Generate submission-ready Word documents for HSR Research Brief.

Outputs:
  - docs/hsr-submission/manuscript.docx (identified, with author info)
  - docs/hsr-submission/manuscript-blinded.docx (anonymous, for peer review)
  - docs/hsr-submission/cover-letter.docx

Format: Times New Roman 12pt, double-spaced, 1" margins, continuous line numbers,
page numbers bottom center, Vancouver-style numbered references.
"""

import os
import re
import sys
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Add metadata cleaner to path
sys.path.insert(0, os.path.expanduser("~/.claude/scripts"))
from clean_docx_metadata import clean_metadata

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_FILE = os.path.join(BASE, "docs", "hsr-research-brief.md")
COVER_FILE = os.path.join(BASE, "docs", "hsr-cover-letter.md")
OUT_DIR = os.path.join(BASE, "docs", "hsr-submission")
os.makedirs(OUT_DIR, exist_ok=True)


# ============================================================
# Helpers
# ============================================================

def set_doc_defaults(doc):
    """Set document-wide defaults: TNR 12pt, 1" margins, double-spaced."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.line_spacing = 2.0  # double-spaced
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Set margins on all sections
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_line_numbers(doc):
    """Add continuous line numbering to document sections via XML."""
    for section in doc.sections:
        sectPr = section._sectPr
        ln_num = parse_xml(
            '<w:lnNumType {} w:countBy="1" w:restart="continuous"/>'.format(
                nsdecls("w")
            )
        )
        sectPr.append(ln_num)


def add_page_numbers(doc):
    """Add page numbers at bottom center of each section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page number field
        run = p.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        fld_char_begin = parse_xml(
            '<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls("w"))
        )
        run._r.append(fld_char_begin)

        run2 = p.add_run()
        instr = parse_xml(
            '<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(
                nsdecls("w")
            )
        )
        run2._r.append(instr)

        run3 = p.add_run()
        fld_char_end = parse_xml(
            '<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls("w"))
        )
        run3._r.append(fld_char_end)


def add_heading(doc, text, level=1):
    """Add a heading with TNR font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.line_spacing = 2.0
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_paragraph(doc, text="", bold=False, italic=False, alignment=None):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p


def add_rich_paragraph(doc, text):
    """Add paragraph with inline formatting: ^N^ for superscript, *text* for italic, **text** for bold."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0

    # Split on formatting markers
    # Handle superscript references ^N^, bold **text**, italic *text*
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\^[0-9,]+\^)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("^") and part.endswith("^"):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
        else:
            run = p.add_run(part)

        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    return p


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)


# ============================================================
# Parse the markdown manuscript
# ============================================================

def parse_manuscript(filepath):
    """Parse the markdown manuscript into structured sections."""
    with open(filepath) as f:
        text = f.read()

    sections = {}

    # Extract title (first # heading)
    title_match = re.search(r'^# (.+)$', text, re.MULTILINE)
    sections["title"] = title_match.group(1) if title_match else ""

    # Extract author block (lines between title and first ---)
    # Use [^\n]+ for title to avoid DOTALL matching across lines
    author_block = re.search(r'^# [^\n]+\n\n(.+?)\n\n---', text, re.DOTALL)
    sections["author_block"] = author_block.group(1).strip() if author_block else ""

    # Extract structured abstract
    abstract_match = re.search(
        r'## Structured Abstract\n\n(.+?)(?=\n---|\n## Introduction)',
        text, re.DOTALL
    )
    sections["abstract"] = abstract_match.group(1).strip() if abstract_match else ""

    # Extract main body sections
    body_sections = [
        "Introduction", "Methods", "Results", "Discussion", "Conclusion"
    ]
    for sec_name in body_sections:
        # Match ## Section through next ## or ---
        pattern = rf'## {sec_name}\n\n(.+?)(?=\n## [A-Z]|\n---\n|\Z)'
        match = re.search(pattern, text, re.DOTALL)
        sections[sec_name.lower()] = match.group(1).strip() if match else ""

    # Extract references
    ref_match = re.search(r'## References\n\n(.+?)(?=\n---|\Z)', text, re.DOTALL)
    sections["references"] = ref_match.group(1).strip() if ref_match else ""

    # Extract exhibit captions (Exhibit sections after references)
    exhibit_matches = re.findall(
        r'## (Exhibit \d+\..+?)\n\n(.+?)(?=\n## |\Z)',
        text, re.DOTALL
    )
    sections["exhibits"] = [(title.strip(), body.strip()) for title, body in exhibit_matches]

    # Extract word count and keywords from author block
    wc_match = re.search(r'\*\*Word count:\*\* (.+)', text)
    sections["word_count"] = wc_match.group(1) if wc_match else ""

    kw_match = re.search(r'\*\*Key words:\*\* (.+)', text)
    sections["keywords"] = kw_match.group(1) if kw_match else ""

    return sections


def write_body_section(doc, section_text, heading_text=None, heading_level=1):
    """Write a body section, handling subsections (### headings) and paragraphs."""
    if heading_text:
        add_heading(doc, heading_text, level=heading_level)

    # Split into subsections and paragraphs
    lines = section_text.split("\n")
    current_para = []

    for line in lines:
        line = line.strip()

        if line.startswith("### "):
            # Flush current paragraph
            if current_para:
                add_rich_paragraph(doc, " ".join(current_para))
                current_para = []
            # Add subheading
            add_heading(doc, line[4:], level=2)

        elif line.startswith("> "):
            # Block quote (used for formula)
            if current_para:
                add_rich_paragraph(doc, " ".join(current_para))
                current_para = []
            p = add_paragraph(doc, line[2:], italic=True)
            p.paragraph_format.left_indent = Inches(0.5)

        elif line.startswith("| "):
            # Table row — collect and render as table
            if current_para:
                add_rich_paragraph(doc, " ".join(current_para))
                current_para = []
            # Tables handled separately in exhibit sections
            # For inline tables, just add as formatted text
            add_paragraph(doc, line)

        elif line == "":
            # Paragraph break
            if current_para:
                add_rich_paragraph(doc, " ".join(current_para))
                current_para = []
        else:
            current_para.append(line)

    # Flush final paragraph
    if current_para:
        add_rich_paragraph(doc, " ".join(current_para))


def write_abstract(doc, abstract_text):
    """Write the structured abstract with bold section labels."""
    add_heading(doc, "Abstract", level=1)

    # Parse abstract sections
    sections = re.split(r'\*\*(\w[\w\s]+?):\*\*', abstract_text)
    # sections[0] is empty or preamble, then alternating: label, content

    for i in range(1, len(sections), 2):
        label = sections[i].strip()
        content = sections[i + 1].strip() if i + 1 < len(sections) else ""

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0

        # Bold label
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.name = "Times New Roman"
        run_label.font.size = Pt(12)

        # Content with inline formatting
        parts = re.split(r'(\^[0-9,]+\^)', content)
        for part in parts:
            if not part:
                continue
            if part.startswith("^") and part.endswith("^"):
                run = p.add_run(part[1:-1])
                run.font.superscript = True
            else:
                run = p.add_run(part)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def write_references(doc, ref_text):
    """Write references section with numbered entries."""
    add_heading(doc, "References", level=1)

    # Split by reference number pattern
    refs = re.split(r'\n(?=\d+\.)', ref_text.strip())

    for ref in refs:
        ref = ref.strip()
        if not ref:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.hanging_indent = Inches(0.5)

        # Parse for italic journal names (*text*)
        parts = re.split(r'(\*[^*]+\*)', ref)
        for part in parts:
            if not part:
                continue
            if part.startswith("*") and part.endswith("*"):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                run = p.add_run(part)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def write_exhibit_tables(doc, sections_data):
    """Write Exhibit 3 and 4 tables on separate pages."""
    # Exhibit 3: Top 10 / Bottom 10 counties table
    for title, body in sections_data["exhibits"]:
        doc.add_page_break()
        add_heading(doc, title, level=1)

        # Check if this exhibit has a table
        table_lines = [l for l in body.split("\n") if l.strip().startswith("|")]

        if table_lines:
            # Parse markdown table
            header_line = table_lines[0]
            # Skip separator line (|---|---|)
            data_lines = [l for l in table_lines[2:] if not re.match(r'^\|[\s-]+\|', l)]

            headers = [cell.strip() for cell in header_line.split("|") if cell.strip()]
            rows = []
            for line in data_lines:
                cells = [cell.strip() for cell in line.split("|") if cell.strip()]
                if cells and cells[0] != "...":
                    rows.append(cells)

            if headers and rows:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Table Grid"

                # Header row
                for i, h in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.bold = True
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(10)

                # Data rows
                for r_idx, row_data in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < len(headers):
                            cell = table.rows[r_idx + 1].cells[c_idx]
                            cell.text = cell_text
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.name = "Times New Roman"
                                    run.font.size = Pt(10)

        # Add note text (non-table lines)
        note_lines = [l for l in body.split("\n")
                      if l.strip() and not l.strip().startswith("|")
                      and not l.strip().startswith("*[See")]
        if note_lines:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 2.0
            run = p.add_run("\n".join(note_lines))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.italic = True


def write_figure_legends(doc, sections_data):
    """Write figure legend page for Exhibits 1 and 2."""
    doc.add_page_break()
    add_heading(doc, "Figure Legends", level=1)

    for title, body in sections_data["exhibits"]:
        if "map" in title.lower() or "Behavioral Health vs" in title:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 2.0

            run = p.add_run(f"{title}")
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

            # Get description text (skip the "[See separate figure]" line)
            desc_lines = [l for l in body.split("\n")
                          if l.strip() and not l.strip().startswith("*[See")]
            if desc_lines:
                run2 = p.add_run("\n" + " ".join(l.strip() for l in desc_lines))
                run2.font.name = "Times New Roman"
                run2.font.size = Pt(12)

            doc.add_paragraph()  # spacing


# ============================================================
# Build documents
# ============================================================

def blind_text(sections_data):
    """Return a deep copy of sections_data with identifying info redacted."""
    import copy
    sd = copy.deepcopy(sections_data)
    # Replace caphegroup.org URL with blinded placeholder
    blind_url = "[URL removed for blinded review]"
    for key in ["abstract", "introduction", "methods", "results", "discussion", "conclusion"]:
        sd[key] = sd[key].replace("caphegroup.org/tools/access-explorer", blind_url)
        sd[key] = sd[key].replace("caphegroup.org", blind_url)
    # Also blind exhibit captions
    sd["exhibits"] = [
        (t, b.replace("caphegroup.org/tools/access-explorer", blind_url)
              .replace("caphegroup.org", blind_url))
        for t, b in sd["exhibits"]
    ]
    return sd


def build_manuscript(sections_data, blinded=False):
    """Build the full manuscript document."""
    if blinded:
        sections_data = blind_text(sections_data)

    doc = Document()
    set_doc_defaults(doc)

    # --- Title page ---
    doc.add_paragraph()  # spacing

    # Title
    p = add_paragraph(doc, sections_data["title"], bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].font.size = Pt(14)

    doc.add_paragraph()  # spacing

    if not blinded:
        # Author info
        author_lines = sections_data["author_block"].split("\n")
        for line in author_lines:
            line = line.strip()
            if not line:
                continue
            # Strip markdown bold
            line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            add_paragraph(doc, line, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        add_paragraph(doc, "[Author information removed for blinded review]",
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

    doc.add_paragraph()

    # Word count
    add_paragraph(doc, f"Word count: {sections_data['word_count']}",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Keywords
    add_paragraph(doc, f"Key words: {sections_data['keywords']}",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Page break after title page
    doc.add_page_break()

    # --- Abstract ---
    write_abstract(doc, sections_data["abstract"])

    doc.add_page_break()

    # --- Main body ---
    # Introduction
    write_body_section(doc, sections_data["introduction"], "Introduction")

    # Methods
    write_body_section(doc, sections_data["methods"], "Methods")

    # Results
    write_body_section(doc, sections_data["results"], "Results")

    # Discussion
    write_body_section(doc, sections_data["discussion"], "Discussion")

    # Conclusion
    write_body_section(doc, sections_data["conclusion"], "Conclusion")

    # Page break before references
    doc.add_page_break()

    # --- References ---
    write_references(doc, sections_data["references"])

    # --- Exhibit tables ---
    write_exhibit_tables(doc, sections_data)

    # --- Figure legends ---
    write_figure_legends(doc, sections_data)

    # --- Page numbers and line numbers ---
    add_page_numbers(doc)
    add_line_numbers(doc)

    # --- Clean metadata ---
    clean_metadata(doc)

    return doc


def build_cover_letter():
    """Build the cover letter document."""
    doc = Document()
    set_doc_defaults(doc)

    with open(COVER_FILE) as f:
        text = f.read()

    # Remove markdown header
    text = re.sub(r'^# Cover Letter\n+', '', text)

    lines = text.split("\n")
    in_body = False

    for line in lines:
        line = line.strip()

        if line == "---":
            # Horizontal rule = spacing
            doc.add_paragraph()
            continue

        if not line:
            if in_body:
                doc.add_paragraph()
            continue

        in_body = True

        # Handle bold markdown
        if line.startswith("**") and ".**" in line:
            # Bold label paragraph like "**Timeliness.**"
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.+?\*\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
        else:
            # Plain or italic text
            line_clean = re.sub(r'\*(.+?)\*', r'\1', line)  # Remove italic markers
            add_paragraph(doc, line_clean)

    # Add page numbers (not line numbers for cover letter)
    add_page_numbers(doc)

    clean_metadata(doc)
    return doc


# ============================================================
# Main
# ============================================================

if __name__ == "__main__":
    import docx.enum.text

    print("Phase 6: Formatting HSR Research Brief for submission...")
    print(f"Source: {MD_FILE}")
    print(f"Output: {OUT_DIR}/")
    print()

    # Parse manuscript
    sections_data = parse_manuscript(MD_FILE)
    print(f"Title: {sections_data['title'][:60]}...")
    print(f"Exhibits found: {len(sections_data['exhibits'])}")

    # Build identified manuscript
    print("\nGenerating manuscript.docx (identified)...")
    doc_id = build_manuscript(sections_data, blinded=False)
    path_id = os.path.join(OUT_DIR, "manuscript.docx")
    doc_id.save(path_id)
    print(f"  Saved: {path_id}")

    # Build blinded manuscript
    print("Generating manuscript-blinded.docx...")
    doc_bl = build_manuscript(sections_data, blinded=True)
    path_bl = os.path.join(OUT_DIR, "manuscript-blinded.docx")
    doc_bl.save(path_bl)
    print(f"  Saved: {path_bl}")

    # Build cover letter
    print("Generating cover-letter.docx...")
    doc_cl = build_cover_letter()
    path_cl = os.path.join(OUT_DIR, "cover-letter.docx")
    doc_cl.save(path_cl)
    print(f"  Saved: {path_cl}")

    # --- Verification ---
    print("\n--- Phase 6 QI Checks ---")
    checks_passed = 0
    checks_total = 0

    # Check 1: Files exist
    for fname in ["manuscript.docx", "manuscript-blinded.docx", "cover-letter.docx"]:
        checks_total += 1
        fpath = os.path.join(OUT_DIR, fname)
        if os.path.exists(fpath):
            size_kb = os.path.getsize(fpath) / 1024
            print(f"  PASS: {fname} exists ({size_kb:.1f} KB)")
            checks_passed += 1
        else:
            print(f"  FAIL: {fname} missing")

    # Check 2: Blinded version has no author name
    checks_total += 1
    doc_check = Document(path_bl)
    full_text = "\n".join(p.text for p in doc_check.paragraphs)
    if "Victoria" not in full_text and "Cholette" not in full_text and "vcholette" not in full_text:
        print("  PASS: Blinded version contains no author identification")
        checks_passed += 1
    else:
        print("  FAIL: Blinded version still contains author identification")
        for i, p in enumerate(doc_check.paragraphs):
            if "Victoria" in p.text or "Cholette" in p.text or "vcholette" in p.text:
                print(f"    Found at paragraph {i}: {p.text[:80]}")

    # Check 3: Identified version has author name
    checks_total += 1
    doc_check_id = Document(path_id)
    full_text_id = "\n".join(p.text for p in doc_check_id.paragraphs)
    if "Victoria Cholette" in full_text_id:
        print("  PASS: Identified version contains author name")
        checks_passed += 1
    else:
        print("  FAIL: Identified version missing author name")

    # Check 4: Structured abstract has all 5 sections
    checks_total += 1
    abstract_sections = ["Objective", "Data Sources", "Study Design",
                         "Principal Findings", "Conclusions"]
    found = [s for s in abstract_sections if s in full_text_id]
    if len(found) == 5:
        print(f"  PASS: Structured abstract has all 5 sections")
        checks_passed += 1
    else:
        missing = [s for s in abstract_sections if s not in full_text_id]
        print(f"  FAIL: Missing abstract sections: {missing}")

    # Check 5: References present and numbered
    checks_total += 1
    ref_count = len(re.findall(r'^\d+\.', full_text_id, re.MULTILINE))
    if ref_count >= 12:
        print(f"  PASS: {ref_count} numbered references found (expected 12)")
        checks_passed += 1
    else:
        print(f"  FAIL: Only {ref_count} numbered references (expected 12)")

    # Check 6: Metadata clean
    checks_total += 1
    cp = doc_check_id.core_properties
    if cp.author == "Victoria Cholette" and "python-docx" not in (cp.comments or ""):
        print(f"  PASS: Metadata clean (author: {cp.author})")
        checks_passed += 1
    else:
        print(f"  FAIL: Metadata not clean (author: {cp.author}, comments: {cp.comments})")

    # Check 7: All 4 exhibits referenced
    checks_total += 1
    exhibits_found = set(re.findall(r'Exhibit [1-4]', full_text_id))
    if len(exhibits_found) >= 4:
        print(f"  PASS: All 4 exhibits referenced ({exhibits_found})")
        checks_passed += 1
    else:
        print(f"  FAIL: Only found {exhibits_found}")

    # Check 8: Line numbers configured
    checks_total += 1
    section = doc_check_id.sections[0]
    sectPr = section._sectPr
    ln_num = sectPr.findall(qn("w:lnNumType"))
    if ln_num:
        print("  PASS: Line numbers configured")
        checks_passed += 1
    else:
        print("  FAIL: Line numbers not configured")

    # Check 9: Double spacing
    checks_total += 1
    normal_style = doc_check_id.styles["Normal"]
    if normal_style.paragraph_format.line_spacing == 2.0:
        print("  PASS: Double-spaced (line_spacing = 2.0)")
        checks_passed += 1
    else:
        print(f"  FAIL: Line spacing = {normal_style.paragraph_format.line_spacing}")

    # Check 10: Font is Times New Roman
    checks_total += 1
    if normal_style.font.name == "Times New Roman":
        print("  PASS: Font is Times New Roman")
        checks_passed += 1
    else:
        print(f"  FAIL: Font is {normal_style.font.name}")

    print(f"\n--- Results: {checks_passed}/{checks_total} PASS ---")
    if checks_passed == checks_total:
        print("Phase 6 COMPLETE: All formatting checks passed.")
    else:
        print(f"Phase 6 INCOMPLETE: {checks_total - checks_passed} failures to address.")
